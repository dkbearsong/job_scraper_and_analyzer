# Libraries
import os
from dotenv import load_dotenv
import json
from docx import Document  # type: ignore
import re
import json
from jobspy import scrape_jobs
import csv
import logging
import random
import time
import numpy as np
import itertools
import yaml
from typing import Any, Dict, List, Optional, Tuple

# Modules
from app.pull_data import DataPuller
from app.text_engine import TextProcessor
from app.ai_engine import AIEngine
from app.archetype_engine import ArchetypeManager, Archetype
from app.llm_classifier import (
    CheapLLMClassifier,
    StrongLLMReranker,
    process_stage_6,
    process_stage_7,
    process_stage_8
)


############################# Global Variables and Configs ############################
load_dotenv()

# Archetype Definitions
with open(os.getenv("ARCHETYPES_CONFIG", ""), 'r') as file:
    ARCHETYPES_CONFIG = json.load(file)

############################## Error Handlers #########################################

def error_logger_crash(error_msg):
    print(error_msg)
    logging.error(error_msg)
    raise ValueError(error_msg)

def error_logger_continue(error_msg):
    print(error_msg)
    logging.error(error_msg)
    return

############################## Helper Functions #######################################
def load_resume_as_text(type):
    # doc = Document(input("Provide the path for the resume file to use: "))
    doc = Document(os.getenv(type))
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text.strip())
    return "\n".join(full_text)

def scrape_single_job_board(new_data, company_url, company=""):
    company_list = [] # list for all jobs in new_data
    for item in new_data['data']:
        if type(item) is not dict:
            continue
        if not item.get('title') or item['title'] == [None] or item['title'] == "":
            continue
        link = item.get('link')
        if link and not link.startswith(("http","https")):
            link = f"{company_url}{link}"
        maker = {
            "company": item['company'] if item.get('company') is not None else company,
            "company_url": company_url,
            "title": item['title'],
            "flexibility": item['flexibility'] if item.get('flexibility') is not None else "NA",
            "url": link,
            "source":new_data['source'] if new_data.get('source') else ""
            }
        if item.get('location') is not None:
            if re.search(r"location", item['location'], re.IGNORECASE):
                item['location'] = re.sub(r'location', "", item['location'], flags=re.IGNORECASE)
            maker["location"] = item['location']
        company_list.append(maker)
    return company_list

def scrape_multi_job_board(new_data, company_url, company):
    full_list = [] # list for all pages combined together
    adjusted_nd = {
        "data":[],
        "status":200,
        "success":True,
        "source":new_data['source']
    }
    for item in new_data['data']:
        adjusted_nd['data'].append(item['jobs'])
    full_list += scrape_single_job_board(adjusted_nd, company_url, company)
    return full_list

async def scrape_sites(i, company_url, dp):
    # print(f"Payload: {i}")
    new_data = await dp.scrape_data(i['strategy'], api_method=i['api_method'])
    new_data['source'] = i['strategy']['source']
    if new_data['status_code'] != 200:
        print(f"Scraping data failed. Error code {new_data['status_code']}. Error: {new_data['error'] if new_data.get('error') is not None else 'No error message provided.'}")
        return
    if 'data' not in new_data:
        print(f"Warning: No 'data' key in response from {i['company']}. Response: {new_data}")
        return
    if len(new_data['data']) != 0:
        try:
            maker = scrape_multi_job_board(new_data, company_url, i['company']) if new_data['data'][0].get('jobs') is not None else scrape_single_job_board(new_data, company_url, i['company'])
        except (KeyError, IndexError, TypeError) as e:
            print(f"Error: {e}\nNew Data: {new_data}")
    else:
        return
    return maker

def scrape_jb(sn, l, rw, ho, **kwargs):
    allowed_keys = {
        "linkedin_fetch_description", 
        "country_indeed", 
        "google_search_term", 
        "search_term"
    }

    filtered_kwargs = {k: v for k, v in kwargs.items() if k in allowed_keys}
    
    jobs = scrape_jobs(
        site_name=[sn],
        location=l,
        results_wanted=rw,
        hours_old=ho,
        **filtered_kwargs
    )
    return jobs

def apply_rule_filters(job: dict, user_preferences: dict) -> bool:
    """
    Evaluates a job against hard constraints.
    Returns True if the job should be SKIPPED, False if it passes.
    """
    job_work_type = job['features'].get('work_type', None)
    if job_work_type:
        job_work_type = job_work_type.lower()
    user_work_types = [t.lower() for t in user_preferences.get('work_types', [])]
    if user_work_types and job_work_type and job_work_type not in user_work_types:
        return True

    job_seniority = job['features'].get('seniority', None)
    if job_seniority:
        job_seniority = job_seniority.lower()
    user_seniority_levels = [s.lower() for s in user_preferences.get('seniority_levels', [])]
    if user_seniority_levels and job_seniority and job_seniority not in user_seniority_levels:
        return True

    job_pay = job['features'].get('pay', "")
    target_pay_range = user_preferences.get('pay_range', '')
    if not job_pay or job_pay == "Not Specified":
        pass
    else:
        if target_pay_range and job_pay:
            try:
                if filter_pay(target_pay_range, job_pay):
                    return True                            
            except (ValueError, IndexError):
                error_logger_continue(f"Error in parsing pay range for job {job['metadata']['job_id']}: {ValueError} at {IndexError}")
                pass

    # 4. Timezone Filter
    job_timezone = job['features'].get('timezone', None)
    if job_timezone:
        job_timezone = job_timezone.lower()
    user_timezones = [tz.lower() for tz in user_preferences.get('timezones', [])]
    
    if user_timezones and job_timezone and job_timezone not in user_timezones:
        return True

    return False

def filter_pay(target_pay_range, job_pay):
    # Parse user's target pay range
    user_min_match = re.search(r'(\d+)(?:k|K)?', target_pay_range)
    user_max_match = re.search(r'-(\d+)(?:k|K)?', target_pay_range)
    if not (user_min_match and user_max_match):
        return None 

    user_min = convert_pay(user_min_match, target_pay_range)
    user_max = convert_pay(user_max_match, target_pay_range)

    job_min, job_max = None, None
    if isinstance(job_pay, str):
        matches = re.findall(r'(\d+)(?:k|K)?', job_pay)
        if matches:
            job_min = int(matches[0]) * 1000 if 'k' in job_pay.lower() else int(matches[0])
        if len(matches) >= 2:
            job_max = int(matches[1]) * 1000 if 'k' in job_pay.lower() else int(matches[1])

    if job_max is not None:
        if job_min is None:
            job_min = job_max
        return True if job_max < user_min or job_min > user_max else False
    if job_min is not None:
        return False if user_min <= job_min <= user_max else True
    return None  # Fallback if no job values could be parsed

def convert_pay(match, context):
    val = int(match.group(1))
    return val * 1000 if 'k' in context.lower() else val

##################################### AI Functions #############################################################################

def call_llm_for_extraction(ai: AIEngine, text: str, provider_name: str | None = None) -> dict:
    """Uses the provided AI engine to extract structured data."""
    if provider_name is None:
        return ai.extract(text)
    return ai.extract(text, provider_name=provider_name)

def generate_embeddings(ai: AIEngine, text: str, provider_name: str | None = None) -> list:
    """Uses the provided AI engine to generate a vector embedding."""
    if provider_name is None:
        return ai.embed(text)
    return ai.embed(text, provider_name=provider_name)

def extract_and_cache_profile(ai: AIEngine, source_path: str, raw_text: str, cache_path: str) -> dict:
    """
    Extracts structured profile data (skills, requirements, summary) from a resume
    or user profile using LLM extraction, with file-modification caching.
    """
    empty_result = {"skills": [], "requirements": [], "summary": ""}

    if not os.path.exists(source_path):
        print(f"Warning: profile source not found at {source_path}")
        return empty_result

    source_mtime = os.path.getmtime(source_path)

    if os.path.exists(cache_path):
        try:
            with open(cache_path, 'r') as f:
                cache = json.load(f)
            cache_time = cache.get('timestamp', 0)
            if cache_time >= source_mtime:
                print(f"Using cached profile from {cache_path}")
                return cache.get('data', empty_result)
        except Exception as e:
            print(f"Warning: failed to read profile cache {cache_path}: {e}")

    if not raw_text:
        print(f"Warning: empty text in {source_path}")
        return empty_result

    print(f"Extracting profile data from {source_path}...")
    ai_data = call_llm_for_extraction(ai, raw_text, provider_name=os.getenv("EXTRACTION_LLM"))
    
    result = empty_result
    if isinstance(ai_data, dict):
        result = {
            "skills": ai_data.get('skills', []),
            "requirements": ai_data.get('requirements', []),
            "summary": ai_data.get('summary', ""),
        }
    elif isinstance(ai_data, str):
        try:
            parsed = json.loads(ai_data)
            if isinstance(parsed, dict):
                result = {
                    "skills": parsed.get('skills', []),
                    "requirements": parsed.get('requirements', []),
                    "summary": parsed.get('summary', ""),
                }
        except Exception:
            pass

    try:
        cache_dir = os.path.dirname(cache_path)
        if cache_dir and not os.path.exists(cache_dir):
            os.makedirs(cache_dir, exist_ok=True)
        with open(cache_path, 'w') as f:
            json.dump({"timestamp": source_mtime, "data": result}, f, indent=2)
        print(f"Cached profile data to {cache_path}")
    except Exception as e:
        print(f"Warning: failed to write profile cache {cache_path}: {e}")

    return result


# =====================================================
# PIPELINE STAGE 0: SETUP AND PROFILE EXTRACTION
# =====================================================

async def pipeline_stage_setup(skip_db: bool = False, verbose: bool = False,
                               resume_text: str = "", profile_text: str = "") -> dict:
    """
    Stage 0: Load environment, resume, profile, create AI engine, extract skills/titles.
    
    Args:
        skip_db: If True, skip database initialization.
        verbose: If True, print detailed debug output.
        resume_text: If provided, use this instead of loading from file.
        profile_text: If provided, use this instead of loading from file.
    
    Returns:
        dict with keys: resume, user_profile, skills, job_titles, dp, tp, ai, user_preferences
    """
    print("=" * 50)
    print("PIPELINE STAGE 0: SETUP")
    print("=" * 50)

    # Load .env variables
    if resume_text:
        resume = resume_text
    else:
        resume = load_resume_as_text("RESUME")
    
    if profile_text:
        user_profile = profile_text
    else:
        user_profile = load_resume_as_text("PROFILE")

    # Create Data Puller Object
    dp = DataPuller(
        dbname=os.getenv("DB_NAME", ""),
        user=os.getenv("DB_USER", ""),
        password=os.getenv("DB_PASSWORD", ""),
        host=os.getenv("DB_HOST", "localhost"),
        port=os.getenv("DB_PORT", "5432")
    )

    # Get sites file from .env
    sites_file = os.getenv("JOB_SITES", "")
    if sites_file and os.path.exists(sites_file):
        sites = dp.load_sites_list(sites_file)
        print("Sites Retrieved")
    else:
        sites = {"name": [], "site": []}
        print("Warning: JOB_SITES file not found or not configured.")

    # Extract skills and job titles from profile
    tp = TextProcessor()
    ai = AIEngine(default_provider_name="lm_studio")
    skills_raw = tp.get_section_content(user_profile, "Skills")
    titles_raw = tp.get_section_content(user_profile, "Job Titles")
    skills = tp.clean_list_from_text(skills_raw)
    job_titles = tp.clean_list_from_text(titles_raw)

    if verbose:
        print(f"--- Profile Extraction ---")
        print(f"Extracted {len(skills)} skills: {skills}")
        print(f"Extracted {len(job_titles)} job titles: {job_titles}")
        print(f"---------------------------\n")

    # Configure logging
    logging.basicConfig(
        filename='app_error.log',
        level=logging.ERROR,
        format='%(asctime)s:%(levelname)s:%(message)s'
    )

    # Load user preferences
    prefs_yaml_path = os.getenv("USER_PREFERENCES_YAML", "")
    if prefs_yaml_path and os.path.exists(prefs_yaml_path):
        with open(prefs_yaml_path, 'r') as f:
            user_preferences = yaml.safe_load(f)
    else:
        user_preferences = {}

    result = {
        "resume": resume,
        "user_profile": user_profile,
        "skills": skills,
        "job_titles": job_titles,
        "dp": dp,
        "tp": tp,
        "ai": ai,
        "user_preferences": user_preferences,
        "sites": sites,
        "sites_file": sites_file,
        "db_config": {
            "dbname": os.getenv("DB_NAME", ""),
            "user": os.getenv("DB_USER", ""),
            "password": os.getenv("DB_PASSWORD", ""),
            "host": os.getenv("DB_HOST", "localhost"),
            "port": os.getenv("DB_PORT", "5432"),
        },
    }

    print("Stage 0 complete: Setup and profile extraction done.")
    return result


# =====================================================
# PIPELINE STAGE 1: SCRAPING
# =====================================================

async def pipeline_stage_scrape(setup_data: dict, skip_db: bool = False,
                                verbose: bool = False) -> List[Dict]:
    """
    Stage 1: Scrape jobs from company career pages and job boards.
    
    Args:
        setup_data: Output from pipeline_stage_setup.
        skip_db: If True, skip database persistence.
        verbose: If True, print detailed debug output.
    
    Returns:
        List[Dict] of scraped jobs in the 'processed_job_pool' format.
    """
    dp = setup_data["dp"]
    tp = setup_data["tp"]
    user_preferences = setup_data.get("user_preferences", {})
    sites = setup_data.get("sites", {"name": [], "site": []})

    print("=" * 50)
    print("PIPELINE STAGE 1: SCRAPING")
    print("=" * 50)

    # ==========================================
    # PART A: Scrape from company career pages
    # ==========================================
    print("--- Company Board Scraping ---")
    site_strategies = []
    data = []

    for i in range(len(sites.get('name', []))):
        strategy_path = f"./site_strategies/{sites['name'][i]}.json"
        if not os.path.exists(strategy_path):
            print(f"Warning: strategy file not found: {strategy_path}")
            continue
        strategy = {
            "company": sites['name'][i],
            "site": sites['site'][i],
            "strategy": dp.load_site_strategies(strategy_path),
            "api_method": "",
        }
        strat = strategy['strategy']
        strategy['api_method'] = (
            "extract-paginated" if strat.get('pagination') is not None
            else ("extract-js" if strat.get("js_config") is not None else "extract")
        )
        if verbose:
            print(f"Company: {strategy['company']} | API method: {strategy['api_method']}")
        site_strategies.append(strategy)
    print(f"Loaded {len(site_strategies)} site strategies.")

    for i in site_strategies:
        company_url = i['strategy'].pop('company_url', None)
        print(f"Scraping {i['company']}...")
        if isinstance(i['strategy']['url'], str):
            d = await scrape_sites(i, company_url, dp)
            if not d:
                continue
            if isinstance(d, list):
                data += d
            else:
                data.append(d)
        elif isinstance(i['strategy']['url'], list):
            for j in i['strategy']['url']:
                new_payload = i
                new_payload['strategy']['url'] = j
                d = await scrape_sites(new_payload, company_url, dp)
                if not d:
                    continue
                if isinstance(d, list):
                    data += d
                else:
                    data.append(d)
    print(f"Total jobs scraped from company boards: {len(data)}")

    if not skip_db:
        dp.load_scraped_data_to_db(data)

    # ==========================================
    # PART B: Scrape from job boards (Indeed, LinkedIn, etc.)
    # ==========================================
    print("--- Job Board Scraping ---")
    job_board_list = ["indeed", "linkedin", "zip_recruiter", "google"]

    search_terms_file = os.getenv("SEARCH_TERMS", "")
    search_terms = []
    if search_terms_file and os.path.exists(search_terms_file):
        try:
            with open(search_terms_file, mode='r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for row in reader:
                    if row:
                        search_terms.append(row[0])
        except Exception as e:
            error_logger_continue(f"Failed to load search terms: {e}")
    if not search_terms:
        print("Warning: No search terms found. Using default.")
        search_terms = ["Software Engineer"]

    target_cities = user_preferences.get('target_cities', [])
    if not target_cities:
        target_cities = ["Remote"]

    jobs = []
    requests_wanted = 200
    delay = {"min": 1, "max": 4}

    for board, st, location in itertools.product(job_board_list, search_terms, target_cities):
        kwa = {}
        if board == "google":
            kwa["google_search_term"] = st
        elif board == "indeed":
            kwa["search_term"] = st
            kwa["country_indeed"] = "USA"
        elif board == "linkedin":
            kwa["search_term"] = st
            kwa["linkedin_fetch_description"] = True
        else:
            kwa["search_term"] = st

        try:
            for i in scrape_jb(board, location, requests_wanted, 24, **kwa):
                if isinstance(i, dict):
                    job = {
                        "source": i.get('site'),
                        "title": i.get('title'),
                        "url": i.get('job_url'),
                        "link": i.get('job_url'),
                        "company": i.get('company'),
                        "pay": f"{i.get('min_amount', '')} - {i.get('max_amount', '')} {i.get('interval', '')}".strip(),
                        "description": i.get('description'),
                        "city": i.get('city'),
                        "state": i.get('state'),
                        "flexibility": i.get('work_type', "NA"),
                        "location": i.get('location', ''),
                    }
                else:
                    job = {
                        "source": None, "title": None, "url": None, "link": None,
                        "company": None, "pay": "", "description": None,
                        "city": None, "state": None, "flexibility": "NA", "location": ""
                    }
                jobs.append(job)
        except Exception as e:
            error_logger_continue(f"Job board scrape failed for {board}/{st}/{location}: {e}")

        time.sleep(random.uniform(delay['min'], delay['max']))

    print(f"Total jobs scraped from job boards: {len(jobs)}")

    if not skip_db:
        dp.load_scraped_data_to_db(jobs)

    # Merge and convert to processed_job_pool format
    all_scraped = data + jobs
    processed_job_pool = []
    seen_ids = set()
    for idx, item in enumerate(all_scraped):
        if not item or not item.get('title'):
            continue
        job_id = item.get('id', idx + 1000)
        if job_id in seen_ids:
            job_id = max(seen_ids) + 1 + idx
        seen_ids.add(job_id)
        processed_job_pool.append({
            "metadata": {
                "job_id": job_id,
                "source": item.get('source', 'scraped'),
            },
            "features": {
                "title": item.get('title', ''),
                "description": item.get('description', ''),
                "pay": item.get('pay', ''),
                "seniority": "NA",
                "work_type": item.get('flexibility', 'NA'),
                "timezone": "NA",
            },
            "embeddings": {
                "description_vector": None,
                "skills_vector": None,
            },
        })

    # If no jobs were scraped (e.g., in test mode), use dummy data as fallback
    if not processed_job_pool:
        from tests.test_data import make_dummy_stage1_output
        print("No real jobs scraped. Using dummy fallback data.")
        processed_job_pool = make_dummy_stage1_output(15)

    print(f"Stage 1 complete: {len(processed_job_pool)} jobs in pool.")
    return processed_job_pool


# =====================================================
# PIPELINE STAGE 2: EMBEDDING GENERATION + LLM EXTRACTION
# =====================================================

async def pipeline_stage_embed_and_extract(jobs: List[Dict], ai_engine: Optional[AIEngine] = None,
                                           text_processor: Optional[TextProcessor] = None,
                                           dp: Optional[DataPuller] = None,
                                           skip_db: bool = False,
                                           verbose: bool = False) -> List[Dict]:
    """
    Stage 2: Run deterministic extraction, LLM extraction, and embedding generation on jobs.
    
    Args:
        jobs: List of job dicts from Stage 1 (or dummy data).
        ai_engine: AIEngine instance for LLM calls and embeddings.
        text_processor: TextProcessor instance for deterministic extraction.
        dp: DataPuller instance for DB operations.
        skip_db: If True, skip database persistence.
        verbose: If True, print detailed debug output.
    
    Returns:
        List[Dict] of jobs with features enriched (skills, requirements, summary, embeddings).
    """
    print("=" * 50)
    print("PIPELINE STAGE 2: EMBEDDING GENERATION + LLM EXTRACTION")
    print("=" * 50)

    if ai_engine is None:
        ai_engine = AIEngine(default_provider_name="lm_studio")
    if text_processor is None:
        text_processor = TextProcessor()

    # First pass: deterministic extraction for jobs that don't have features yet
    print(f"Starting deterministic extraction on {len(jobs)} jobs...")
    processed_job_pool = []
    for idx, raw_job in enumerate(jobs):
        if raw_job is None:
            continue
        
        # If the job is already in the processed format (from previous stage), use it
        if 'features' in raw_job and 'metadata' in raw_job:
            processed_job_pool.append(raw_job)
            continue

        # Build from raw format
        job_id = raw_job.get('id') or raw_job.get('metadata', {}).get('job_id', idx + 1)
        description = raw_job.get('description', raw_job.get('job_summary', ''))
        title = raw_job.get('title', raw_job.get('job_name', ''))
        work_type = raw_job.get('flexibility', 'NA')
        pay = raw_job.get('pay', raw_job.get('pay_range', ''))

        if not description or len(description) < 50:
            if verbose:
                error_logger_continue(f"Warning: insufficient description for job ID {job_id}")
            continue

        extracted_data = {
            "metadata": {
                "job_id": job_id,
                "source": "scraped",
            },
            "features": {
                "title": title,
                "description": description,
                "pay": pay if pay else text_processor.extract_salary(description),
                "seniority": text_processor.detect_seniority(description),
                "work_type": work_type if work_type and work_type != "NA" else text_processor.detect_work_type(description),
                "timezone": text_processor.detect_timezone(description),
            },
            "embeddings": {
                "description_vector": None,
                "skills_vector": None,
            },
        }
        processed_job_pool.append(extracted_data)

    print(f"Successfully extracted data for {len(processed_job_pool)} jobs.")

    # Second pass: LLM extraction + embedding generation
    print(f"Starting AI/LLM Pass on {len(processed_job_pool)} jobs...")
    for index, job in enumerate(processed_job_pool):
        if verbose:
            print(f"Processing job {index + 1}/{len(processed_job_pool)}: {job['features']['title']}")

        if not job or 'features' not in job:
            error_logger_continue(f"Warning: job at index {index} has invalid structure")
            continue

        description = job['features'].get('description', '')
        if not description:
            error_logger_continue(f"Warning: job at index {index} has no description")
            continue

        # Extract skills, requirements and summary via LLM
        ai_data = call_llm_for_extraction(ai_engine, description, provider_name=os.getenv("EXTRACTION_LLM", "lm_studio"))

        skills = []
        requirements = []
        summary = ""

        if isinstance(ai_data, dict):
            skills = ai_data.get('skills', [])
            requirements = ai_data.get('requirements', [])
            summary = ai_data.get('summary', "")
        elif isinstance(ai_data, str):
            try:
                parsed = json.loads(ai_data)
                if isinstance(parsed, dict):
                    skills = parsed.get('skills', [])
                    requirements = parsed.get('requirements', [])
                    summary = parsed.get('summary', "")
            except Exception:
                pass

        features = job['features']
        features['skills'] = skills or []
        features['requirements'] = requirements or []
        features['summary'] = summary or ""

        # Vector generation
        title_text = features['title']
        if title_text:
            job['embeddings']['title_vector'] = generate_embeddings(ai_engine, title_text, provider_name=os.getenv("EMBEDDINGS_LLM", "lm_studio"))
        else:
            job['embeddings']['title_vector'] = []

        skills_text = ", ".join(features['skills'])
        if skills_text:
            job['embeddings']['skills_vector'] = generate_embeddings(ai_engine, skills_text, provider_name=os.getenv("EMBEDDINGS_LLM", "lm_studio"))
        else:
            job['embeddings']['skills_vector'] = []

        requirements_text = ", ".join(features['requirements'])
        if requirements_text:
            job['embeddings']['requirements_vector'] = generate_embeddings(ai_engine, requirements_text, provider_name=os.getenv("EMBEDDINGS_LLM", "lm_studio"))
        else:
            job['embeddings']['requirements_vector'] = []

        summary_text = features['summary']
        if summary_text:
            job['embeddings']['description_vector'] = generate_embeddings(ai_engine, summary_text, provider_name=os.getenv("EMBEDDINGS_LLM", "lm_studio"))
        else:
            job['embeddings']['description_vector'] = []

        time.sleep(1)

    print("AI/LLM Pass Complete.")

    # Persist to database if not skipping
    if dp and not skip_db:
        # Save embeddings
        print("Saving embeddings to database...")
        embedding_updates = []
        for job in processed_job_pool:
            emb = job.get('embeddings', {})
            embedding_updates.append({
                "job_id": job['metadata']['job_id'],
                "title_embedding": emb.get('title_vector'),
                "skills_embedding": emb.get('skills_vector'),
                "responsibilities_embedding": emb.get('requirements_vector'),
                "description_embedding": emb.get('description_vector'),
            })
        if embedding_updates:
            dp.save_job_embeddings(embedding_updates)
            print(f"Saved {len(embedding_updates)} jobs' embeddings.")

        # Save metadata
        print("Updating job records in database with metadata...")
        job_updates = []
        for job in processed_job_pool:
            job_updates.append({
                "id": job['metadata']['job_id'],
                "pay_range": job['features'].get('pay'),
                "seniority": job['features'].get('seniority'),
                "work_type": job['features'].get('work_type'),
                "timezone": job['features'].get('timezone'),
            })
        dp.update_job_metadata(job_updates)
        print(f"Updated {len(job_updates)} job records.")

    print(f"Stage 2 complete: {len(processed_job_pool)} jobs processed.")
    return processed_job_pool


# =====================================================
# PIPELINE STAGE 3: RULE FILTERING
# =====================================================

async def pipeline_stage_rule_filter(jobs: List[Dict], user_preferences: Optional[dict] = None,
                                     dp: Optional[DataPuller] = None,
                                     skip_db: bool = False,
                                     verbose: bool = False) -> List[Dict]:
    """
    Stage 3: Apply hard-constraint rule filtering to the job pool.
    
    Args:
        jobs: List of job dicts from Stage 2.
        user_preferences: Dict of user preferences (work types, seniority, pay, timezones).
        dp: DataPuller instance for DB sync.
        skip_db: If True, skip database persistence.
        verbose: If True, print detailed debug output.
    
    Returns:
        List[Dict] of all jobs (processed_job_pool) with 'skip' boolean added,
        plus the 'active_job_pool' key if returning the full context.
    """
    print("=" * 50)
    print("PIPELINE STAGE 3: RULE FILTERING")
    print("=" * 50)

    if user_preferences is None:
        user_preferences = {}

    print(f"Starting Rule-Based Filtering on {len(jobs)} jobs...")

    for job in jobs:
        job['skip'] = apply_rule_filters(job, user_preferences)

    skipped_count = sum(1 for j in jobs if j['skip'])
    print(f"Filtering complete. {skipped_count} jobs skipped, {len(jobs) - skipped_count} active.")

    # Sync skip status to database
    if dp and not skip_db:
        print("Syncing skip status to database...")
        skipped_ids = [job['metadata']['job_id'] for job in jobs if job.get('skip')]
        if skipped_ids:
            dp.bulk_update_skip_status(skipped_ids)
            print(f"Updated {len(skipped_ids)} jobs as skipped.")
        else:
            print("No jobs to skip in database.")

    print(f"Stage 3 complete. Active jobs: {len(jobs) - skipped_count}")
    return jobs


# =====================================================
# PIPELINE STAGE 4: ARCHETYPE ENGINE INTEGRATION
# =====================================================

async def pipeline_stage_archetype_integration(jobs: List[Dict], ai_engine: Optional[AIEngine] = None,
                                                dp: Optional[DataPuller] = None,
                                                setup_data: Optional[dict] = None,
                                                skip_db: bool = False,
                                                verbose: bool = False) -> Tuple[List[Dict], ArchetypeManager]:
    """
    Stage 4: Load/integrate archetypes (benchmarks + user profile/resume).
    
    Args:
        jobs: List of job dicts from Stage 3 (with skip flags).
        ai_engine: AIEngine for generating archetype embeddings.
        dp: DataPuller for DB operations.
        setup_data: Setup data containing resume/user_profile text.
        skip_db: If True, skip DB operations.
        verbose: If True, print detailed output.
    
    Returns:
        Tuple of (active_jobs, archetype_manager).
    """
    print("=" * 50)
    print("PIPELINE STAGE 4: ARCHETYPE ENGINE INTEGRATION")
    print("=" * 50)

    if ai_engine is None:
        ai_engine = AIEngine(default_provider_name="lm_studio")

    archetype_manager = ArchetypeManager()

    # Get active (non-skipped) jobs
    active_jobs = [j for j in jobs if not j.get('skip')]
    print(f"Active jobs for archetype comparison: {len(active_jobs)}")

    print("Synchronizing candidate archetypes and benchmarks...")

    # Extract structured profile data using LLM with file-modification caching
    resume_text = ""
    profile_text = ""
    if setup_data:
        resume_text = setup_data.get("resume", "")
        profile_text = setup_data.get("user_profile", "")

    resume_data = extract_and_cache_profile(
        ai_engine,
        os.getenv("RESUME", "") if not resume_text else "archetype_profiles/resume_cache.json",
        resume_text if resume_text else "No resume text provided.",
        "archetype_profiles/resume_cache.json"
    )
    profile_data = extract_and_cache_profile(
        ai_engine,
        os.getenv("PROFILE", "") if not profile_text else "archetype_profiles/user_profile_cache.json",
        profile_text if profile_text else "No profile text provided.",
        "archetype_profiles/user_profile_cache.json"
    )

    all_archetypes = ARCHETYPES_CONFIG + [
        {
            "name": "Resume",
            "title": ", ".join(resume_data.get("skills", [])),
            "skills": "\n".join(resume_data.get("skills", [])),
            "responsibilities": "\n".join(resume_data.get("requirements", [])),
            "summary": resume_data.get("summary", ""),
            "type": "resume"
        },
        {
            "name": "User Profile",
            "title": ", ".join(profile_data.get("skills", [])),
            "skills": "\n".join(profile_data.get("skills", [])),
            "responsibilities": "\n".join(profile_data.get("requirements", [])),
            "summary": profile_data.get("summary", ""),
            "type": "user_profile"
        }
    ]

    for arch_config in all_archetypes:
        cached_arch = None
        if dp:
            cached_arch = dp.get_archetype_embeddings(arch_config['name'])

        if cached_arch:
            archetype_manager.add_archetype(Archetype(
                name=arch_config['name'],
                type=cached_arch['archetype_type'],
                title_embedding=np.array(cached_arch['title_embedding']),
                skills_embedding=np.array(cached_arch['skills_embedding']),
                responsibilities_embedding=np.array(cached_arch['responsibilities_embedding']),
                metadata=cached_arch.get('metadata', {})
            ))
        else:
            print(f"Generating new embeddings for archetype: {arch_config['name']}")
            new_arch = archetype_manager.load_archetype(
                name=arch_config['name'],
                archetype_data=arch_config,
                archetype_type=arch_config.get("type", "benchmark")
            )
            # Persist to DB for future runs
            if dp and not skip_db:
                dp.save_archetype_embeddings({
                    "archetype_name": new_arch.name,
                    "archetype_type": new_arch.type,
                    "title_embedding": new_arch.title_embedding.tolist() if new_arch.title_embedding is not None else None,
                    "skills_embedding": new_arch.skills_embedding.tolist() if new_arch.skills_embedding is not None else None,
                    "responsibilities_embedding": new_arch.responsibilities_embedding.tolist() if new_arch.responsibilities_embedding is not None else None,
                    "metadata": json.dumps(new_arch.metadata)
                })

    print(f"Stage 4 complete: Loaded {len(archetype_manager.archetypes)} archetypes.")
    return active_jobs, archetype_manager


# =====================================================
# PIPELINE STAGE 5: VECTOR SCORING WITH ARCHETYPES
# =====================================================

async def pipeline_stage_vector_scoring(jobs: List[Dict], archetype_manager: Optional[ArchetypeManager] = None,
                                         dp: Optional[DataPuller] = None,
                                         skip_db: bool = False,
                                         verbose: bool = False) -> List[Dict]:
    """
    Stage 5: Compare jobs to archetypes, compute semantic scores, and filter by threshold.
    
    Args:
        jobs: List of active job dicts from Stage 4.
        archetype_manager: ArchetypeManager with loaded archetypes.
        dp: DataPuller for DB persistence.
        skip_db: If True, skip DB operations.
        verbose: If True, print detailed output.
    
    Returns:
        List[Dict] of filtered jobs with semantic scores and archetype matches.
    """
    print("=" * 50)
    print("PIPELINE STAGE 5: VECTOR SCORING WITH ARCHETYPES")
    print("=" * 50)

    if archetype_manager is None:
        archetype_manager = ArchetypeManager()

    print("Comparing jobs to archetypes...")
    for index, job in enumerate(jobs):
        if verbose:
            print(f"Processing job {index + 1}/{len(jobs)}: {job['features']['title']}")

        matches = archetype_manager.compare_job_to_archetypes(job)
        job['retrieval_metadata'] = archetype_manager.generate_retrieval_metadata(job, matches)
        job['archetype_matches'] = matches

    print("Archetype comparison complete.")

    # Import adjustment functions from vector_engine
    from app.vector_engine import apply_keyword_adjustments, apply_metadata_adjustments

    print("Applying weighted semantic scoring...")
    for job in jobs:
        matches = job.get('archetype_matches', [])
        if not matches:
            continue

        best_match = matches[0]

        title_similarity = best_match.get('title_similarity', 0.0)
        skills_similarity = best_match.get('skills_similarity', 0.0)
        responsibility_similarity = best_match.get('responsibility_similarity', 0.0)

        # Weighted semantic score
        semantic_score = (
            0.40 * title_similarity +
            0.35 * skills_similarity +
            0.25 * responsibility_similarity
        )

        # Keyword adjustments
        job_skills = job.get('features', {}).get('skills', [])
        job_title = job.get('features', {}).get('title', '')
        semantic_score = apply_keyword_adjustments(semantic_score, job_skills, job_title)

        # Metadata adjustments
        job_meta = {
            "is_remote": job.get('features', {}).get('work_type', '').lower() == 'remote',
            "salary": 0,
            "days_old": 30
        }
        pay_range = job.get('features', {}).get('pay', '')
        if pay_range:
            import re as _re
            numbers = _re.findall(r'\d+(?:,\d+)?', pay_range.replace(',', ''))
            if len(numbers) >= 2:
                job_meta["salary"] = int(numbers[1])
            elif len(numbers) == 1:
                job_meta["salary"] = int(numbers[0])

        semantic_score = apply_metadata_adjustments(semantic_score, job_meta)

        # Clamp and normalize
        semantic_score = max(0.0, min(1.0, semantic_score))
        score_percent = int(round(semantic_score * 100))

        job['semantic_score'] = semantic_score
        job['semantic_score_percent'] = score_percent
        job['title_similarity'] = title_similarity
        job['skills_similarity'] = skills_similarity
        job['responsibility_similarity'] = responsibility_similarity
        job['adjusted_score'] = semantic_score
        job['best_archetype'] = best_match.get('archetype_name', '')

    # Generate retrieval metadata
    for job in jobs:
        if 'retrieval_metadata' not in job:
            job['retrieval_metadata'] = {}
        if 'semantic_score' in job:
            job['retrieval_metadata']['semantic_score'] = job['semantic_score']
            job['retrieval_metadata']['semantic_score_percent'] = job['semantic_score_percent']
            job['retrieval_metadata']['best_archetype'] = job.get('best_archetype', '')

    # Filter by threshold
    MIN_SCORE_THRESHOLD = 0.72
    TARGET_COUNT = 20

    filtered_job_pool = []
    for job in jobs:
        if job.get('semantic_score', 0) >= MIN_SCORE_THRESHOLD:
            filtered_job_pool.append(job)

    # Fallback: add top-X if not enough jobs meet threshold
    if len(filtered_job_pool) < TARGET_COUNT:
        sorted_jobs = sorted(jobs, key=lambda x: x.get('semantic_score', 0), reverse=True)
        top_n = max(TARGET_COUNT - len(filtered_job_pool), 1)
        for job in sorted_jobs[:top_n]:
            if job not in filtered_job_pool:
                filtered_job_pool.append(job)

    print(f"Filtered to {len(filtered_job_pool)} jobs (threshold >= {MIN_SCORE_THRESHOLD}).")

    # Persist vector scores
    if dp and not skip_db:
        try:
            dp.save_vector_scores(filtered_job_pool)
            print(f"Persisted {len(filtered_job_pool)} vector scores to database.")
        except Exception as e:
            error_logger_continue(f"Vector score persistence failed: {e}")

    # Print top jobs
    sorted_filtered = sorted(filtered_job_pool, key=lambda x: x.get('semantic_score', 0), reverse=True)
    print("\n=== Semantic Score-Based Job Shortlisting ===")
    for i, job in enumerate(sorted_filtered[:5]):
        print(f"{i+1}. {job['features']['title']}")
        print(f"   Score: {job.get('semantic_score_percent', 'N/A')}% | Archetype: {job.get('best_archetype', 'None')}")

    print(f"Stage 5 complete: {len(filtered_job_pool)} jobs in filtered pool.")
    return filtered_job_pool


# =====================================================
# PIPELINE STAGE 6: CHEAP LLM CLASSIFICATION
# =====================================================

async def pipeline_stage_cheap_llm(jobs: List[Dict], setup_data: Optional[dict] = None,
                                    dp: Optional[DataPuller] = None,
                                    skip_db: bool = False,
                                    verbose: bool = False) -> List[Dict]:
    """
    Stage 6: Run cheap (fast) LLM classification on the filtered job pool.
    
    Args:
        jobs: List of filtered job dicts from Stage 5.
        setup_data: Setup data containing user_profile and skills.
        dp: DataPuller for DB persistence.
        skip_db: If True, skip DB operations.
        verbose: If True, print detailed output.
    
    Returns:
        List[Dict] of shortlisted jobs with 'cheap_llm_result' added.
    """
    print("=" * 50)
    print("PIPELINE STAGE 6: CHEAP LLM CLASSIFICATION")
    print("=" * 50)

    user_profile = setup_data.get("user_profile", "") if setup_data else ""
    skills = setup_data.get("skills", []) if setup_data else []

    if not user_profile:
        print("Warning: No user profile available for classification.")

    # Initialize cheap LLM classifier
    cheap_llm_provider = os.getenv("CHEAP_LLM_PROVIDER", "gemini")
    cheap_classifier = CheapLLMClassifier(provider=cheap_llm_provider)

    # Run Stage 6 on filtered job pool
    shortlisted_jobs = await process_stage_6(
        jobs=jobs,
        classifier=cheap_classifier,
        candidate_profile=user_profile,
        candidate_skills=skills,
        batch_size=5
    )

    # Persist results
    if dp and not skip_db:
        print("Persisting Stage 6 results to database...")
        try:
            dp.save_cheap_llm_results(shortlisted_jobs)
            print(f"Persisted {len(shortlisted_jobs)} cheap LLM results.")
        except Exception as e:
            error_logger_continue(f"Stage 6 persistence failed: {e}")

    print(f"Stage 6 complete: {len(shortlisted_jobs)} jobs shortlisted.")
    return shortlisted_jobs


# =====================================================
# PIPELINE STAGE 7: STRONG LLM RERANKING
# =====================================================

async def pipeline_stage_strong_llm(jobs: List[Dict], setup_data: Optional[dict] = None,
                                     dp: Optional[DataPuller] = None,
                                     skip_db: bool = False,
                                     verbose: bool = False) -> List[Dict]:
    """
    Stage 7: Run strong (deep) LLM reranking on top candidates from Stage 6.
    
    Args:
        jobs: List of shortlisted job dicts from Stage 6.
        setup_data: Setup data containing user_profile and skills.
        dp: DataPuller for DB persistence.
        skip_db: If True, skip DB operations.
        verbose: If True, print detailed output.
    
    Returns:
        List[Dict] of deeply analyzed jobs with 'strong_llm_result' added.
    """
    print("=" * 50)
    print("PIPELINE STAGE 7: STRONG LLM RERANKING")
    print("=" * 50)

    user_profile = setup_data.get("user_profile", "") if setup_data else ""
    skills = setup_data.get("skills", []) if setup_data else []

    # Initialize strong LLM reranker
    strong_llm_provider = os.getenv("STRONG_LLM_PROVIDER", "claude")
    strong_reranker = StrongLLMReranker(provider=strong_llm_provider)

    # Configure how many jobs to deeply analyze
    top_n_for_deep_analysis = int(os.getenv("TOP_N_DEEP_ANALYSIS", "15"))

    # Run Stage 7 on top candidates from Stage 6
    deeply_analyzed_jobs = await process_stage_7(
        jobs=jobs,
        reranker=strong_reranker,
        candidate_profile=user_profile,
        candidate_skills=skills,
        top_n=top_n_for_deep_analysis
    )

    # Persist results
    if dp and not skip_db:
        print("Persisting Stage 7 results to database...")
        try:
            dp.save_strong_llm_results(deeply_analyzed_jobs)
            print(f"Persisted {len(deeply_analyzed_jobs)} strong LLM results.")
        except Exception as e:
            error_logger_continue(f"Stage 7 persistence failed: {e}")

    print(f"Stage 7 complete: {len(deeply_analyzed_jobs)} jobs deeply analyzed.")
    return deeply_analyzed_jobs


# =====================================================
# PIPELINE STAGE 8: FINAL APPLICATION QUEUE
# =====================================================

async def pipeline_stage_final_queue(jobs: List[Dict], dp: Optional[DataPuller] = None,
                                      skip_db: bool = False,
                                      verbose: bool = False) -> List[Dict]:
    """
    Stage 8: Generate the final ranked application queue from deeply analyzed jobs.
    
    Args:
        jobs: List of deeply analyzed job dicts from Stage 7.
        dp: DataPuller for DB persistence.
        skip_db: If True, skip DB operations.
        verbose: If True, print detailed output.
    
    Returns:
        List[Dict] of ranked jobs with final_score, priority, apply_recommendation.
    """
    print("=" * 50)
    print("PIPELINE STAGE 8: FINAL APPLICATION QUEUE")
    print("=" * 50)

    # Run Stage 8 to generate final ranked queue
    final_queue = await process_stage_8(jobs=jobs)

    # Persist results
    if dp and not skip_db:
        print("Persisting final application queue to database...")
        try:
            dp.save_final_queue(final_queue)
            print(f"Persisted {len(final_queue)} jobs to final queue.")
        except Exception as e:
            error_logger_continue(f"Final queue persistence failed: {e}")

    # Print detailed final queue
    print("\n" + "=" * 60)
    print("DETAILED FINAL APPLICATION QUEUE")
    print("=" * 60)

    for i, job in enumerate(final_queue[:10], 1):
        features = job.get('features', {})
        cheap_result = job.get('cheap_llm_result', {})
        strong_result = job.get('strong_llm_result', {})

        print(f"\n{i}. {features.get('title', 'Unknown')}")
        print(f"   Priority: {job.get('priority', 'unknown').upper()}")
        print(f"   Final Score: {job.get('final_score', 0):.1f}/100")
        print(f"   Recommendation: {job.get('apply_recommendation', 'maybe').upper()}")
        print(f"   Cheap LLM Fit: {cheap_result.get('fit_score', 0)}/100")
        print(f"   Strong LLM Score: {strong_result.get('final_score', 0)}/100")

    print(f"\nStage 8 complete: {len(final_queue)} jobs in final queue.")
    return final_queue


# =====================================================
# ORIGINAL MAIN (calls pipeline stages)
# =====================================================

async def main():
    """
    Original main function - runs the full pipeline end-to-end.
    """
    print("=" * 60)
    print("FULL JOB SCRAPING AND ANALYSIS PIPELINE")
    print("=" * 60)

    # Stage 0: Setup
    setup_data = await pipeline_stage_setup(verbose=False)
    dp = setup_data["dp"]
    ai = setup_data["ai"]
    tp = setup_data["tp"]

    # Stage 1: Scrape
    processed_job_pool = await pipeline_stage_scrape(
        setup_data=setup_data,
        skip_db=False,
        verbose=False,
    )

    # Stage 2: Embedding Generation + LLM Extraction
    processed_job_pool = await pipeline_stage_embed_and_extract(
        jobs=processed_job_pool,
        ai_engine=ai,
        text_processor=tp,
        dp=dp,
        skip_db=False,
        verbose=False,
    )

    # Stage 3: Rule Filtering
    processed_job_pool = await pipeline_stage_rule_filter(
        jobs=processed_job_pool,
        user_preferences=setup_data.get("user_preferences", {}),
        dp=dp,
        skip_db=False,
    )

    # Stage 4: Archetype Engine Integration
    active_jobs, archetype_manager = await pipeline_stage_archetype_integration(
        jobs=processed_job_pool,
        ai_engine=ai,
        dp=dp,
        setup_data=setup_data,
        skip_db=False,
    )

    # Stage 5: Vector Scoring
    filtered_job_pool = await pipeline_stage_vector_scoring(
        jobs=active_jobs,
        archetype_manager=archetype_manager,
        dp=dp,
        skip_db=False,
    )

    # Stage 6: Cheap LLM Classification
    shortlisted_jobs = await pipeline_stage_cheap_llm(
        jobs=filtered_job_pool,
        setup_data=setup_data,
        dp=dp,
        skip_db=False,
    )

    # Stage 7: Strong LLM Reranking
    deeply_analyzed_jobs = await pipeline_stage_strong_llm(
        jobs=shortlisted_jobs,
        setup_data=setup_data,
        dp=dp,
        skip_db=False,
    )

    # Stage 8: Final Application Queue
    final_queue = await pipeline_stage_final_queue(
        jobs=deeply_analyzed_jobs,
        dp=dp,
        skip_db=False,
    )

    # Pipeline Summary
    print("\n" + "=" * 60)
    print("PIPELINE COMPLETE")
    print("=" * 60)
    print(f"Total jobs processed: {len(processed_job_pool)}")
    print(f"Jobs after semantic filtering: {len(filtered_job_pool)}")
    print(f"Jobs after cheap LLM classification: {len(shortlisted_jobs)}")
    print(f"Jobs after strong LLM reranking: {len(deeply_analyzed_jobs)}")
    print(f"Final application queue: {len(final_queue)} jobs")

    return final_queue


# Entry point
if __name__ == "__main__":
    import asyncio
    asyncio.run(main())